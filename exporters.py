# CONFIDENTIAL - TRADE SECRET
# Proprietary KorumOS source code. Access is limited to authorized personnel
# and collaborators operating under written confidentiality obligations.

import base64
import csv
import io
import json
import os
import re
import tempfile
import zipfile
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from xml.sax.saxutils import escape
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- SHARED UTILS ---

def _timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def _safe_filename_part(s):
    if not s: return "UNTITLED"
    return re.sub(r'[^\w\-]', '_', str(s))[:30]

def _output_path(filename, output_dir=None):
    if not output_dir:
        output_dir = os.path.join(os.getcwd(), "exports")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    return os.path.join(output_dir, filename)

def _as_text(val):
    if val is None: return ""
    return str(val).strip()

def _flatten_structured_value(val):
    """Convert dicts/lists from synthesis JSON into readable prose text."""
    if val is None:
        return ""
    if isinstance(val, str):
        return val
    if isinstance(val, list):
        lines = []
        for item in val:
            if isinstance(item, dict):
                # Scenario-style dicts: {name, description, timeline, implication}
                parts = []
                name = item.get('name', '')
                if name:
                    parts.append(f"**{name}**")
                for k in ('description', 'timeline', 'implication'):
                    v = item.get(k)
                    if v:
                        parts.append(f"{k.title()}: {v}")
                # Catch any other keys
                for k, v in item.items():
                    if k not in ('name', 'description', 'timeline', 'implication') and v:
                        parts.append(f"{k.replace('_', ' ').title()}: {v}")
                lines.append("\n".join(parts))
            else:
                lines.append(f"- {_as_text(item)}")
        return "\n\n".join(lines) if any(isinstance(i, dict) for i in val) else "\n".join(lines)
    if isinstance(val, dict):
        lines = []
        for k, v in val.items():
            label = k.replace('_', ' ').upper()
            if isinstance(v, list):
                lines.append(f"**{label}**")
                for item in v:
                    if isinstance(item, dict):
                        lines.append(f"- {_flatten_structured_value(item)}")
                    else:
                        lines.append(f"- {_as_text(item)}")
            elif isinstance(v, dict):
                lines.append(f"**{label}**")
                lines.append(_flatten_structured_value(v))
            elif isinstance(v, str):
                lines.append(f"**{label}**: {v}")
            else:
                lines.append(f"**{label}**: {_as_text(v)}")
        return "\n".join(lines)
    return str(val)


# Phase titles for anonymized pull quote attributions (replaces provider names)
_PHASE_TITLES = ["ANALYST", "ARCHITECT", "CRITIC", "INTEGRATOR", "COMPOSER"]

# Canonical section order for RESEARCH reports (matches WORKFLOW_DNA output_structure)
_RESEARCH_SECTION_ORDER = [
    "executive_summary", "key_signals", "system_context", "scenario_analysis",
    "critical_challenges", "risks", "tradeoffs", "decision", "action_priorities",
    "evidence",
    "execution_considerations", "confidence_assessment", "confidence",
    "final_assessment",
]

def _reorder_sections(sections, order=None):
    """Reorder sections dict to match canonical output_structure.
    Sections not in the order list are appended at the end."""
    if not order:
        order = _RESEARCH_SECTION_ORDER
    ordered = []
    remaining = dict(sections)
    for key in order:
        if key in remaining:
            ordered.append((key, remaining.pop(key)))
    # Append any sections not in the canonical order
    for key, val in remaining.items():
        ordered.append((key, val))
    return ordered



def _packet_backed_risks_mode(data, sections):
    """
    True when this export is packet-backed and has normalized risk content.
    """
    packet = (data or {}).get("decision_packet") or {}
    risks_text = str((sections or {}).get("risks") or "").strip()
    return bool(packet) and bool(risks_text)

def _filter_sections(section_items, data, sections_dict):
    pb_mode = _packet_backed_risks_mode(data, sections_dict)
    filtered = []
    for sid, content in section_items:
        key = str(sid).strip().lower()
        if key == "critical_challenges" and pb_mode:
            continue
        if key == "risks" and not pb_mode:
            continue
        filtered.append((sid, content))
    return filtered

def _extract_parts(o):
    meta = o.get("meta") or {}
    sections = o.get("sections") or {}
    structured = o.get("structured_data") or {}
    interrogations = o.get("interrogations") or []
    verifications = o.get("verifications") or []
    decision_packet = o.get("decision_packet") or {}
    if decision_packet:
        meta, sections, structured, interrogations, verifications = _packet_preferred_parts(
            decision_packet,
            meta,
            sections,
            structured,
            interrogations,
            verifications,
        )
    return meta, sections, structured, interrogations, verifications


def _packet_list(items):
    return items if isinstance(items, list) else []


def _packet_dict(value):
    return value if isinstance(value, dict) else {}


def _packet_text_list(items, field_name=None):
    values = []
    for item in _packet_list(items):
        if isinstance(item, dict):
            if field_name:
                text = _as_text(item.get(field_name))
            else:
                text = _as_text(item.get("text") or item.get("value") or item.get("title") or item.get("point"))
        else:
            text = _as_text(item)
        if text:
            values.append(text)
    return values


def _prefer_packet_value(packet_value, fallback_value=""):
    packet_text = _as_text(packet_value)
    if packet_text:
        return packet_value
    return fallback_value


def _packet_score(packet, fallback_meta):
    confidence = _packet_dict(packet.get("confidence"))
    raw_score = confidence.get("score")
    if raw_score is not None and raw_score != "":
        try:
            return _normalize_truth_score(raw_score)
        except Exception:
            return None

    if not confidence:
        fallback_score = fallback_meta.get("truth_score") or fallback_meta.get("composite_truth_score")
        if fallback_score is not None and fallback_score != "":
            return _normalize_truth_score(fallback_score)
    return None


def _packet_status(packet, score):
    confidence = _packet_dict(packet.get("confidence"))
    raw_band = _as_text(confidence.get("band")).upper()
    if raw_band in ("HIGH", "MEDIUM", "LOW"):
        if raw_band == "HIGH":
            return "Recommended" if score >= 80 else "Conditional"
        if raw_band == "MEDIUM":
            return "Conditional"
        return "Fail"
    if score is not None:
        return _confidence_status_from_score(score)
    return ""


def _packet_timeline_bucket(timeline_text):
    value = _as_text(timeline_text).lower()
    if not value:
        return "Immediate"
    if "mid" in value:
        return "Mid-Term"
    if "near" in value or "30" in value or "90" in value or "short" in value:
        return "Near-Term"
    if "immediate" in value or "now" in value or "week" in value or "day" in value:
        return "Immediate"
    return "Immediate"


def _packet_action_priorities_text(actions):
    grouped = {"Immediate": [], "Near-Term": [], "Mid-Term": []}
    for item in _packet_list(actions):
        if isinstance(item, dict):
            action = _as_text(item.get("action"))
            owner = _as_text(item.get("owner"))
            if not action:
                continue
            line = action if not owner else f"{action} ({owner})"
            bucket = _packet_timeline_bucket(item.get("timeline"))
            grouped.setdefault(bucket, []).append(line)
        else:
            text = _as_text(item)
            if text:
                grouped["Immediate"].append(text)

    lines = []
    for heading in ("Immediate", "Near-Term", "Mid-Term"):
        items = grouped.get(heading) or []
        if not items:
            continue
        lines.append(f"**{heading.upper()}**")
        lines.extend(f"- {item}" for item in items)
    return "\n".join(lines).strip()


def _packet_confidence_text(packet, score, status):
    confidence = _packet_dict(packet.get("confidence"))
    basis = _as_text(confidence.get("basis"))
    assumptions = _packet_text_list(packet.get("assumptions"))
    unknowns = _packet_text_list(packet.get("unknowns"))

    lines = []
    if score is not None:
        lines.append(f"Score: {score} / 100")
    if status:
        lines.append(f"Status: {status}")
    if basis:
        if lines:
            lines.append("")
        lines.extend(["Key Drivers:", f"- {basis}"])
    if assumptions:
        if lines:
            lines.append("")
        lines.append("Assumptions:")
        lines.extend(f"- {item}" for item in assumptions[:5])
    if unknowns:
        if lines:
            lines.append("")
        lines.append("Limitations:")
        lines.extend(f"- {item}" for item in unknowns[:5])
    return "\n".join(lines)


def _packet_risk_text(packet):
    risks = []
    for item in _packet_list(packet.get("risk_vectors")):
        if isinstance(item, dict):
            title = _as_text(item.get("title"))
            description = _as_text(item.get("description"))
            mitigation = _as_text(item.get("mitigation"))
            line = f"- {title}"
            details = []
            if description:
                details.append(description)
            if mitigation:
                details.append(f"Mitigation: {mitigation}")
            if details:
                line = f"{line} {'. '.join(details)}"
            risks.append(line)
        else:
            text = _as_text(item)
            if text:
                risks.append(f"- {text}")
    return "\n".join(risks)


def _packet_evidence_text(packet):
    lines = []
    for item in _packet_list(packet.get("evidence_trace")):
        if isinstance(item, dict):
            point = _as_text(item.get("point"))
            support = _as_text(item.get("support"))
            if point and support:
                lines.append(f"- {point}: {support}")
            elif point:
                lines.append(f"- {point}")
        else:
            text = _as_text(item)
            if text:
                lines.append(f"- {text}")
    return "\n".join(lines)


def _packet_tradeoffs_text(packet):
    lines = []
    for item in _packet_list(packet.get("alternatives_rejected")):
        if isinstance(item, dict):
            option = _as_text(item.get("option"))
            reason = _as_text(item.get("reason_rejected"))
            if option and reason:
                lines.append(f"- {option}: {reason}")
            elif option:
                lines.append(f"- {option}")
        else:
            text = _as_text(item)
            if text:
                lines.append(f"- {text}")
    return "\n".join(lines)


def _packet_execution_text(packet):
    assumptions = _packet_text_list(packet.get("assumptions"))
    unknowns = _packet_text_list(packet.get("unknowns"))
    parts = []
    if assumptions:
        parts.append("**ASSUMPTIONS**")
        parts.extend(f"- {item}" for item in assumptions[:5])
    if unknowns:
        if parts:
            parts.append("")
        parts.append("**UNKNOWNS**")
        parts.extend(f"- {item}" for item in unknowns[:5])
    return "\n".join(parts)


def _packet_key_signals_text(packet):
    claims = []
    for item in _packet_list(packet.get("verified_claims"))[:4]:
        if isinstance(item, dict):
            claim = _as_text(item.get("claim"))
            if claim:
                claims.append(f"- {claim}")
        else:
            text = _as_text(item)
            if text:
                claims.append(f"- {text}")
    return "\n".join(claims)


def _packet_verification_blocks(packet):
    blocks = []
    for item in _packet_list(packet.get("verified_claims")):
        if not isinstance(item, dict):
            continue
        claim = _as_text(item.get("claim"))
        if not claim:
            continue
        source_ref = _as_text(item.get("source_ref"))
        raw_status = _as_text(item.get("status")).upper() or "UNVERIFIED"
        if raw_status in ("VERIFIED", "ACCURATE"):
            status = "VERIFIED"
            verdict = "ACCURATE"
        elif raw_status in ("FALSE", "FLAGGED"):
            status = "FLAGGED"
            verdict = "INACCURATE"
        else:
            status = "CONDITIONAL"
            verdict = "CONDITIONAL"
        blocks.append({
            "type": "evidence",
            "label": "SOURCE VERIFICATION",
            "status": status,
            "claim": claim,
            "verdict": verdict,
            "evidence_points": [claim],
            "sources": [source_ref] if source_ref else [],
            "challenges": [],
            "text": claim,
        })
    return blocks


def _packet_preferred_parts(packet, fallback_meta, fallback_sections, fallback_structured, fallback_interrogations, fallback_verifications):
    export_meta = _packet_dict(packet.get("export_metadata"))
    score = _packet_score(packet, fallback_meta)
    status = _packet_status(packet, score)
    title = _as_text(_prefer_packet_value(packet.get("decision_headline"), fallback_meta.get("title")))
    summary = _as_text(_prefer_packet_value(packet.get("executive_summary"), fallback_meta.get("summary")))
    go_no_go = _packet_dict(packet.get("go_no_go_call"))
    rationale = _as_text(go_no_go.get("rationale"))
    decision_value = _as_text(go_no_go.get("decision"))

    meta = dict(fallback_meta or {})
    meta["title"] = title or meta.get("title")
    meta["summary"] = summary or meta.get("summary")
    meta["workflow"] = _as_text(
        _prefer_packet_value(
            export_meta.get("workflow") or export_meta.get("report_type"),
            meta.get("workflow") or "RESEARCH",
        )
    )
    meta["generated_at"] = _as_text(_prefer_packet_value(export_meta.get("generated_at"), meta.get("generated_at")))
    meta["_export_packet_primary"] = True
    if score is not None:
        meta["composite_truth_score"] = score
        meta["truth_score"] = score
        meta["confidence_score"] = score
    else:
        meta.pop("composite_truth_score", None)
        meta.pop("truth_score", None)
        meta.pop("confidence_score", None)
    if status:
        meta["status"] = status
    else:
        meta.pop("status", None)

    packet_risks = _packet_risk_text(packet)
    sections = dict(fallback_sections or {})
    sections.update({
        "executive_summary": summary or sections.get("executive_summary", ""),
        "key_signals": _packet_key_signals_text(packet) or sections.get("key_signals", ""),
        "critical_challenges": sections.get("critical_challenges", "") if packet_risks else sections.get("critical_challenges", ""),
        "risks": packet_risks or sections.get("risks", ""),
        "tradeoffs": _packet_tradeoffs_text(packet) or sections.get("tradeoffs", ""),
        "decision": "\n".join(filter(None, [f"Decision: {decision_value}" if decision_value else "", rationale])) or sections.get("decision", ""),
        "action_priorities": _packet_action_priorities_text(packet.get("immediate_actions")) or sections.get("action_priorities", ""),
        "evidence": _packet_evidence_text(packet) or sections.get("evidence", ""),
        "execution_considerations": _packet_execution_text(packet) or sections.get("execution_considerations", ""),
        "confidence_assessment": _packet_confidence_text(packet, score, status) or sections.get("confidence_assessment", ""),
        "final_assessment": title or sections.get("final_assessment", ""),
    })

    structured = dict(fallback_structured or {})
    if score is not None:
        key_metrics = [{"metric": "Decision Score", "value": f"{score} / 100"}]
    else:
        key_metrics = []
    structured["key_metrics"] = key_metrics
    structured["action_items"] = structured.get("action_items") or _packet_list(packet.get("immediate_actions"))
    structured["risks"] = structured.get("risks") or _packet_list(packet.get("risk_vectors"))

    verifications = fallback_verifications or _packet_verification_blocks(packet)
    interrogations = fallback_interrogations or []
    return meta, sections, structured, interrogations, verifications

def _artifact_label(art):
    label = _as_text(art.get('title') or art.get('name') or '').strip()
    raw_content = _as_text(art.get('content', '')).lstrip()
    if raw_content.startswith('<svg'):
        return 'CHART'
    if not label:
        art_type = _as_text(art.get('type') or '').upper()
        if art_type in ('VISUALIZATION', 'MERMAID', 'CHART', 'SVG'):
            label = 'CHART'
        elif art_type in ('CSV', 'TABLE', 'DATA'):
            label = 'DATA TABLE'
        else:
            # Only use recognized types — suppress stray labels like CODE, TEXT, etc.
            label = ''
    return label.upper()


def _is_raw_svg_artifact(art):
    return _as_text((art or {}).get('content', '')).lstrip().startswith('<svg')

def _report_artifacts(o):
    """Filter snippets for inclusion in the report."""
    apps = o.get("docked_snippets") or []
    return [a for a in apps if a.get('includeInReport') is True]


def _normalize_theme(theme_id):
    theme_key = _as_text(theme_id).upper() or "BONE_FIELD"
    return theme_key if theme_key in THEME_COLORS else "BONE_FIELD"


def _safe_paragraph(text, style):
    return Paragraph(escape(_as_text(text)).replace("\n", "<br/>"), style)


def _artifact_unavailable_note(styles):
    return Paragraph("Chart artifact unavailable for export capture.", styles['ExecAudit'])


def _closing_stamp_parts(session_id):
    timestamp = datetime.now().strftime("%B %d, %Y %I:%M %p")
    return (
        "Produced by KorumOS Decision Intelligence",
        "Proprietary \u00b7 Confidential",
        f"SESSION {session_id} \u00b7 {timestamp}",
    )


def _pdf_pipe_join(*parts):
    cleaned = [_as_text(part).strip() for part in parts if _as_text(part).strip()]
    return " | ".join(cleaned)


def _pdf_status_label(status):
    status_lower = _as_text(status).lower()
    if status_lower in ("flagged", "challenged", "inaccurate"):
        return "FLAGGED"
    if status_lower == "verified" or status_lower == "accurate":
        return "VERIFIED"
    return "CONDITIONAL"


def _clean_cell_text(text):
    t = _as_text(text)
    t = re.sub(r"\[/?METRIC_ANCHOR\]", "", t)
    t = re.sub(r"^#{1,6}\s*", "", t, flags=re.MULTILINE)       # markdown headers
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)                     # bold (paired)
    t = re.sub(r"\*\*", "", t)                                   # stray unclosed bold markers
    t = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", t)  # italic
    t = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", t)            # markdown links
    t = re.sub(r"\[([A-Z_ ]+)\]", "", t)                       # signal tags like [AMBER]
    return t.strip()


def _table_from_rows(headers, rows):
    if not headers or not rows:
        return None
    clean_headers = [_clean_cell_text(h) for h in headers]
    clean_rows = [[_clean_cell_text(cell) for cell in row] for row in rows if any(_as_text(cell) for cell in row)]
    if not clean_rows:
        return None
    return {"headers": clean_headers, "rows": clean_rows}


def _parse_structured_table_payload(raw_payload):
    try:
        data = json.loads(raw_payload.strip())
    except Exception:
        return None

    headers = []
    rows = []
    if isinstance(data, dict):
        if isinstance(data.get("headers"), list):
            headers = [_as_text(h) for h in data.get("headers", [])]
        if isinstance(data.get("rows"), list):
            rows = data["rows"]
        elif isinstance(data.get("data"), list):
            rows = data["data"]
        elif all(isinstance(v, (str, int, float)) for v in data.values()):
            headers = ["Field", "Value"]
            rows = [[k, v] for k, v in data.items()]
    elif isinstance(data, list):
        rows = data

    if rows and isinstance(rows[0], dict):
        if not headers:
            headers = list(rows[0].keys())
        matrix = [[row.get(h, "") for h in headers] for row in rows]
        return _table_from_rows(headers, matrix)

    if rows and isinstance(rows[0], (list, tuple)):
        if not headers and len(rows) > 1:
            headers = [_as_text(cell) for cell in rows[0]]
            rows = rows[1:]
        elif not headers:
            headers = [f"Column {idx + 1}" for idx in range(len(rows[0]))]
        return _table_from_rows(headers, rows)

    return None


def _parse_markdown_table_block(block):
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if len(lines) < 2 or not all(line.startswith("|") and line.endswith("|") for line in lines[:2]):
        return None

    separator = re.sub(r"[|\-:\s]", "", lines[1])
    if separator:
        return None

    headers = [cell.strip() for cell in lines[0].strip("|").split("|")]
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines[2:]]
    return _table_from_rows(headers, rows)


def _parse_delimited_table_block(block):
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if len(lines) < 2:
        return None

    delimiter = None
    for candidate in ("\t", ","):
        if all(candidate in line for line in lines[:2]):
            delimiter = candidate
            break
    if not delimiter:
        return None

    parsed = [[cell.strip() for cell in line.split(delimiter)] for line in lines]
    width = len(parsed[0])
    if width < 2 or any(len(row) != width for row in parsed):
        return None
    return _table_from_rows(parsed[0], parsed[1:])


def _parse_key_value_block(lines):
    pairs = []
    for line in lines:
        match = re.match(r"^([^:]{2,120}):\s+(.+)$", line)
        if not match:
            return None
        pairs.append([match.group(1).strip(), match.group(2).strip()])
    if len(pairs) < 2:
        return None
    return _table_from_rows(["Metric", "Value"], pairs)


def _extract_content_blocks(text):
    raw_text = _as_text(text)
    if not raw_text:
        return []

    blocks = []
    last_end = 0
    pattern = re.compile(r"\[STRUCTURED_TABLE\]([\s\S]*?)\[/STRUCTURED_TABLE\]", re.IGNORECASE)
    for match in pattern.finditer(raw_text):
        if match.start() > last_end:
            blocks.extend(_extract_content_blocks(raw_text[last_end:match.start()]))
        table_block = _parse_structured_table_payload(match.group(1))
        if table_block:
            blocks.append({"type": "table", **table_block})
        last_end = match.end()

    if last_end:
        if last_end < len(raw_text):
            blocks.extend(_extract_content_blocks(raw_text[last_end:]))
        return blocks

    # Pre-pass: extract full evidence blocks (SOURCE VERIFICATION / INTERROGATION)
    # These span multiple paragraphs — capture from header to next header or end
    _EV_BLOCK = re.compile(
        r"((?:SOURCE VERIFICATION|VERIFICATION|INTERROGATION)\s*\[[^\]]+\][\s\S]*?)(?=(?:SOURCE VERIFICATION|VERIFICATION|INTERROGATION)\s*\[|$)",
        re.IGNORECASE
    )
    ev_regions = []
    for ev_match in _EV_BLOCK.finditer(raw_text):
        ev_text = ev_match.group(1).strip()
        if len(ev_text) > 30:  # skip trivially short matches
            ev_regions.append((ev_match.start(), ev_match.end(), ev_text))

    # If we found evidence blocks, extract them and process remaining text separately
    if ev_regions:
        cursor = 0
        for start, end, ev_text in ev_regions:
            if start > cursor:
                # Process non-evidence text before this block
                blocks.extend(_extract_content_blocks(raw_text[cursor:start]))
            blocks.append(_parse_evidence_block(ev_text))
            cursor = end
        if cursor < len(raw_text):
            blocks.extend(_extract_content_blocks(raw_text[cursor:]))
        return blocks

    for chunk in re.split(r"\n\s*\n", raw_text):
        chunk = chunk.strip()
        if not chunk:
            continue

        markdown_table = _parse_markdown_table_block(chunk)
        if markdown_table:
            blocks.append({"type": "table", **markdown_table})
            continue

        delimited_table = _parse_delimited_table_block(chunk)
        if delimited_table:
            blocks.append({"type": "table", **delimited_table})
            continue

        lines = [line.strip() for line in chunk.splitlines() if line.strip()]
        if len(lines) >= 3 and ":" not in lines[0]:
            kv_table = _parse_key_value_block(lines[1:])
            if kv_table:
                blocks.append({"type": "heading", "text": lines[0]})
                blocks.append({"type": "table", **kv_table})
                continue

        kv_table = _parse_key_value_block(lines)
        if kv_table:
            blocks.append({"type": "table", **kv_table})
            continue

        blocks.append({"type": "paragraph", "text": chunk})

    return blocks


def _parse_evidence_block(text):
    """Parse a SOURCE VERIFICATION / INTERROGATION block into structured fields."""
    lines = text.split("\n")

    # Extract header status
    header_match = re.match(
        r"(SOURCE VERIFICATION|VERIFICATION|INTERROGATION)\s*\[([^\]]+)\]",
        lines[0].strip(), re.IGNORECASE
    )
    status = header_match.group(2).strip().upper() if header_match else "UNKNOWN"
    label = header_match.group(1).strip().upper() if header_match else "VERIFICATION"

    claim = ""
    verdict = ""
    evidence_points = []
    sources = []
    challenges = []
    current_section = "body"  # body, evidence, sources

    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        clean_text = stripped.replace("**", "").replace("*", "").strip()
        clean_lower = clean_text.lower()

        # Claim line
        if stripped.lower().startswith("claim:"):
            claim_text = stripped[6:].strip().strip('"').strip("'")
            claim = claim_text
            continue

        # Verdict line: **ACCURATE.** or **INACCURATE.**
        if (
            stripped.startswith("**")
            and stripped.endswith("**")
            and len(stripped) < 40
            and ":" not in clean_text
            and re.match(
                r"^(ACCURATE|INACCURATE|PARTIALLY ACCURATE|MIXED|CONDITIONAL|FLAGGED|UNCERTAIN)\.?$",
                clean_text,
                re.IGNORECASE,
            )
        ):
            verdict = clean_text.rstrip(".")
            continue

        # Section markers — must look like a header line ("Sources:", "Authoritative references:")
        source_header = re.match(r"^sources?\s*:\s*(.*)", clean_lower)
        if "authoritative reference" in clean_lower or source_header:
            current_section = "sources"
            # If there's inline content after "Source: <value>", capture it
            if source_header:
                inline_val = clean_text[source_header.start(1):].strip()
                if inline_val:
                    sources.append(inline_val)
            continue
        if "challenged by" in clean_lower or "challenge:" in clean_lower:
            challenge_text = clean_text.lstrip("⚑").strip()
            challenges.append(challenge_text)
            continue

        # Bullet points
        if stripped.startswith("- "):
            clean = stripped[2:].strip().replace("**", "")
            if current_section == "sources":
                sources.append(clean)
            else:
                evidence_points.append(clean)
                current_section = "evidence"
            continue

        # Continuation text
        if current_section == "sources":
            sources.append(clean_text)
        elif current_section == "evidence":
            evidence_points.append(clean_text)

    return {
        "type": "evidence",
        "status": status,
        "label": label,
        "claim": claim,
        "verdict": verdict,
        "evidence_points": evidence_points,
        "sources": sources,
        "challenges": challenges,
        "text": text,  # raw fallback
    }


def _first_table_block(text):
    for block in _extract_content_blocks(text):
        if block.get("type") == "table":
            return block
    return None


def _decode_image_data(data_url, max_width=500, max_height=300):
    """Decode a base64 data URL into a ReportLab Image flowable."""
    try:
        if not data_url or not data_url.startswith('data:image/'):
            return None
        header, encoded = data_url.split(',', 1)
        img_bytes = base64.b64decode(encoded)
        buf = io.BytesIO(img_bytes)
        img = Image(buf)
        # Scale to fit within max bounds while preserving aspect ratio
        w, h = img.drawWidth, img.drawHeight
        if w > max_width:
            scale = max_width / w
            w, h = w * scale, h * scale
        if h > max_height:
            scale = max_height / h
            w, h = w * scale, h * scale
        img.drawWidth = w
        img.drawHeight = h
        return img
    except Exception:
        return None


def _decode_image_bytes(data_url, max_width_inches=4.5, max_height_inches=3.5):
    """Decode a base64 data URL into (BytesIO, width_inches, height_inches) for Word embedding.
    Scales to fit within max bounds while preserving aspect ratio."""
    try:
        if not data_url or not data_url.startswith('data:image/'):
            return None, None, None
        _, encoded = data_url.split(',', 1)
        img_bytes = base64.b64decode(encoded)

        # Read PNG dimensions from header (bytes 16-23 for width/height)
        w_px, h_px = 800, 400  # fallback
        if len(img_bytes) > 24 and img_bytes[:4] == b'\x89PNG':
            import struct
            w_px = struct.unpack('>I', img_bytes[16:20])[0]
            h_px = struct.unpack('>I', img_bytes[20:24])[0]

        # Convert to inches at 96 DPI and scale to fit
        w_in = w_px / 96.0
        h_in = h_px / 96.0
        if w_in > max_width_inches:
            scale = max_width_inches / w_in
            w_in, h_in = w_in * scale, h_in * scale
        if h_in > max_height_inches:
            scale = max_height_inches / h_in
            w_in, h_in = w_in * scale, h_in * scale

        return io.BytesIO(img_bytes), w_in, h_in
    except Exception:
        return None, None, None


def _artifact_matches(art, section_title, provider_name):
    """Check if an artifact belongs next to a given node by matching content keywords."""
    label = _artifact_label(art).upper()
    content = _as_text(art.get('content', '')).upper()[:500]
    sec = section_title.upper()
    prov = provider_name.upper()
    # Match if artifact label/content mentions the provider or section topic
    return (prov and prov != 'KORUM' and (prov in label or prov in content)) or \
           (sec and len(sec) > 3 and (sec in label or sec in content))


def _assign_artifacts_to_nodes(artifacts, section_items, contributors):
    """Map each artifact to its FIRST matching node. Each artifact renders exactly once."""
    assigned = {i: [] for i in range(len(section_items))}
    claimed = set()  # track artifact indices already assigned
    for i, (sid, _content) in enumerate(section_items):
        sec_title = sid.replace("_", " ").upper()
        prov_name = "KORUM"
        if i < len(contributors):
            prov_name = _as_text(contributors[i].get('provider', '')).upper()
        for ai, a in enumerate(artifacts):
            if ai not in claimed and _artifact_matches(a, sec_title, prov_name):
                assigned[i].append(a)
                claimed.add(ai)
    # Unmatched artifacts distribute across nodes in artifact order so they stay near
    # the report body instead of piling up as an appendix on the last node.
    if section_items:
        next_idx = 0
        for ai, a in enumerate(artifacts):
            if ai in claimed:
                continue
            assigned[next_idx].append(a)
            next_idx = (next_idx + 1) % len(section_items)
    return assigned


def _pdf_col_widths(total_width, column_count):
    if column_count <= 1:
        return [total_width]
    if column_count == 2:
        return [total_width * 0.64, total_width * 0.36]
    return [total_width / column_count] * column_count


def _build_pdf_table(block, styles, total_width, palette):
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers or not rows:
        return None

    table_rows = [
        [Paragraph(f"<b>{escape(_as_text(cell))}</b>", styles["ExecTableHead"]) for cell in headers]
    ]
    table_rows.extend(
        [
            [Paragraph(escape(_as_text(cell)).replace("\n", "<br/>"), styles["ExecTableCell"]) for cell in row]
            for row in rows
        ]
    )
    tbl = Table(table_rows, colWidths=_pdf_col_widths(total_width, len(headers)), repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(palette["shade_alt"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(palette["accent_dark"])),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(palette["line"])),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(palette["line_strong"])),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(palette["paper"]), colors.HexColor(palette["bg_shade"])]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    return tbl


def _pull_quote_color(status):
    status_lower = _as_text(status).lower()
    if status_lower in ("flagged", "challenged"):
        return SEM_RED, "#FDF2F2"
    if status_lower == "verified":
        return SEM_GREEN, "#F2F8F4"
    return SEM_BLUE, "#F0F3FA"


def _build_pull_quote(claim_text, status, provider, role, session_id, styles):
    border_color, bg_color = _pull_quote_color(status)
    quote_style = ParagraphStyle(
        'PQ', parent=styles['ExecBody'], fontSize=8, leading=11,
        fontName='Helvetica-Oblique', textColor=colors.HexColor("#2f2b24"),
    )
    attr_style = ParagraphStyle(
        'PQAttr', parent=styles['ExecAudit'], fontSize=6.5,
        fontName='Courier', textColor=colors.HexColor(SEM_MUTED),
    )
    role_part = f" | {role}" if role else ""
    content = [
        Paragraph(f"&ldquo;{escape(_as_text(claim_text))}&rdquo;", quote_style),
        Spacer(1, 4),
        Paragraph(f"- {provider}{role_part} | {session_id}", attr_style),
    ]
    tbl = Table([[content]], colWidths=[520])
    tbl.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(bg_color)),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('LINEBEFORE', (0,0), (0,-1), 3, colors.HexColor(border_color)),
    ]))
    return tbl


def _render_pdf_blocks(blocks, styles, total_width, palette):
    flowables = []
    for block in blocks:
        if block["type"] == "paragraph":
            flowables.append(_safe_paragraph(block["text"], styles["ExecBody"]))
            flowables.append(Spacer(1, 8))
        elif block["type"] == "heading":
            flowables.append(_safe_paragraph(block["text"], styles["ExecSectionSubhead"]))
            flowables.append(Spacer(1, 5))
        elif block["type"] == "table":
            table = _build_pdf_table(block, styles, total_width, palette)
            if table:
                flowables.append(table)
                flowables.append(Spacer(1, 10))
        elif block["type"] == "evidence":
            ev_flowables = _build_pdf_evidence_block(block, styles, total_width, palette)
            flowables.extend(ev_flowables)
    return flowables


def _build_pdf_evidence_block(block, styles, total_width, palette):
    """KORUM-style evidence block: claim strip + status + evidence + provenance."""
    block_label = _as_text(block.get("label") or "SOURCE VERIFICATION").upper()
    status = block.get("status", "").upper()
    claim = block.get("claim", "")
    evidence_points = block.get("evidence_points", [])
    sources = block.get("sources", [])
    challenges = block.get("challenges", [])

    if status in ("INACCURATE", "FLAGGED", "CHALLENGED"):
        border_color = colors.HexColor("#A45A52")
        status_label = "FLAGGED"
        bg_color = colors.HexColor("#FDF6F5")
    elif status in ("ACCURATE", "VERIFIED"):
        border_color = colors.HexColor("#5B7F5E")
        status_label = "VERIFIED"
        bg_color = colors.HexColor("#F4F8F5")
    else:
        border_color = colors.HexColor("#4A6A7A")
        status_label = "CONDITIONAL"
        bg_color = colors.HexColor("#F0F3FA")

    flowables = [Spacer(1, 6)]

    # ── 1. CLAIM STRIP (left-bordered) ───────────────────────────────
    if claim:
        claim_parts = [
            Paragraph(f"<b><i>{escape(claim)}</i></b>", styles["ExecBody"]),
            Spacer(1, 3),
            Paragraph(f"- {status_label} CLAIM", styles["ExecAudit"]),
        ]
        claim_tab = Table([[claim_parts]], colWidths=[total_width - 30])
        claim_tab.setStyle(TableStyle([
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LINEBEFORE', (0, 0), (0, -1), 3, border_color),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        flowables.append(claim_tab)
        flowables.append(Spacer(1, 6))

    # ── 2. STATUS STRIP ──────────────────────────────────────────────
    status_text = f"{block_label} | {status_label}"
    status_tab = Table([[Paragraph(f"<b>{escape(status_text)}</b>", styles["ExecAudit"])]], colWidths=[total_width - 30])
    status_tab.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), bg_color),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    flowables.append(status_tab)
    flowables.append(Spacer(1, 6))

    # ── 3. CHALLENGES ────────────────────────────────────────────────
    for challenge in challenges[:2]:
        ch_parts = [Paragraph(f"<i>{escape(challenge)}</i>", styles["ExecAudit"])]
        ch_tab = Table([[ch_parts]], colWidths=[total_width - 30])
        ch_tab.setStyle(TableStyle([
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LINEBEFORE', (0, 0), (0, -1), 2, colors.HexColor("#A45A52")),
        ]))
        flowables.append(ch_tab)
        flowables.append(Spacer(1, 4))

    # ── 4. WHY THIS HOLDS ────────────────────────────────────────────
    if evidence_points:
        flowables.append(Paragraph(f"<b>WHY THIS HOLDS:</b>", styles["ExecAudit"]))
        flowables.append(Spacer(1, 3))
        for point in evidence_points[:5]:
            flowables.append(Paragraph(f"- {escape(point)}", styles["ExecBody"]))
            flowables.append(Spacer(1, 2))
        flowables.append(Spacer(1, 4))

    # ── 5. PROVENANCE ────────────────────────────────────────────────
    if sources:
        source_summary = " | ".join(s[:60] for s in sources[:3])
        challenge_status = "PASSED" if status in ("ACCURATE", "VERIFIED") else "FAILED" if status in ("INACCURATE", "FLAGGED") else "PENDING"
        prov_parts = [
            Paragraph(f"SOURCE: {escape(source_summary)}", styles["ExecAudit"]),
            Spacer(1, 2),
            Paragraph(f"<b>STATUS: {status_label} | CHALLENGE: {challenge_status}</b>", styles["ExecAudit"]),
        ]
        prov_tab = Table([[prov_parts]], colWidths=[total_width - 30])
        prov_tab.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), bg_color),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        flowables.append(prov_tab)

    flowables.append(Spacer(1, 12))
    return flowables


def _resolve_session_id(meta, intelligence_object):
    mission_ctx = intelligence_object.get("_mission_context") or {}
    return _as_text(
        meta.get("session_id")
        or meta.get("id")
        or intelligence_object.get("session_id")
        or intelligence_object.get("id")
        or mission_ctx.get("session_id")
        or "KO-INT-9999"
    ).upper()


def _normalize_truth_score(raw_score):
    if raw_score is None or raw_score == "":
        return 85
    try:
        score = float(raw_score)
    except (TypeError, ValueError):
        return 85
    if score <= 1:
        score *= 100
    return int(round(score))


def _confidence_status_from_score(score):
    if score >= 80:
        return "Recommended"
    if score >= 70:
        return "Conditional"
    return "Fail"


def _extract_labeled_segment(text, label, stop_labels):
    if not text:
        return ""
    stop_pattern = "|".join(re.escape(lbl) for lbl in stop_labels)
    match = re.search(
        rf"{re.escape(label)}\s*:\s*(.*?)(?=(?:{stop_pattern})\s*:|$)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return match.group(1).strip(" .\n") if match else ""


def _split_export_items(text):
    if not text:
        return []
    normalized = text.replace("\u2022", ";").replace("•", ";")
    parts = [re.sub(r"^\-\s*", "", item).strip(" .") for item in re.split(r"\s*;\s*|\n+", normalized)]
    return [item for item in parts if item]


def _score_from_confidence_text(text, meta=None):
    match = re.search(r"(\d{1,3})\s*/\s*100", _as_text(text))
    if match:
        return _normalize_truth_score(match.group(1))

    score_sources = []
    if meta:
        score_sources.extend([
            meta.get("truth_score"),
            meta.get("score"),
            meta.get("confidence_score"),
        ])
    for raw in score_sources:
        if raw is not None and raw != "":
            return _normalize_truth_score(raw)

    lowered = _as_text(text).lower()
    mappings = [
        ("very high", 88),
        ("moderate-high", 78),
        ("moderate high", 78),
        ("high", 82),
        ("moderate", 74),
        ("low", 65),
    ]
    for phrase, score in mappings:
        if phrase in lowered:
            return score
    return 82


def _normalize_action_priorities_text(text):
    normalized = _as_text(text)
    replacements = [
        (r"(\*\*)?\s*IMMEDIATE\s*\([^)]+\)\s*(\*\*)?", "**IMMEDIATE**"),
        (r"(\*\*)?\s*NEAR[\-\s]?TERM\s*\([^)]+\)\s*(\*\*)?", "**NEAR-TERM**"),
        (r"(\*\*)?\s*MID[\-\s]?TERM\s*\([^)]+\)\s*(\*\*)?", "**MID-TERM**"),
        (r"(\*\*)?\s*SHORT[\-\s]?TERM\s*\([^)]+\)\s*(\*\*)?", "**NEAR-TERM**"),
    ]
    for pattern, replacement in replacements:
        normalized = re.sub(pattern, replacement, normalized, flags=re.IGNORECASE)
    return normalized


def _normalize_confidence_text(text, meta=None):
    raw = _as_text(text).strip()
    if not raw:
        return raw

    score = _score_from_confidence_text(raw, meta)
    status = _confidence_status_from_score(score)
    assumptions = _split_export_items(
        _extract_labeled_segment(
            raw,
            "KEY ASSUMPTIONS",
            ["LIMITATIONS", "KEY DRIVERS", "STATUS", "SCORE"],
        )
    )
    limitations = _split_export_items(
        _extract_labeled_segment(
            raw,
            "LIMITATIONS",
            ["KEY ASSUMPTIONS", "KEY DRIVERS", "STATUS", "SCORE"],
        )
    )
    key_drivers = _split_export_items(
        _extract_labeled_segment(
            raw,
            "KEY DRIVERS",
            ["LIMITATIONS", "KEY ASSUMPTIONS", "STATUS", "SCORE"],
        )
    )

    lines = [
        f"Score: {score} / 100",
        f"Status: {status}",
    ]

    if key_drivers:
        lines.extend(["", "Key Drivers:"])
        lines.extend(f"- {item}" for item in key_drivers[:5])

    if assumptions:
        lines.extend(["", "Assumptions:"])
        lines.extend(f"- {item}" for item in assumptions[:5])

    if limitations:
        lines.extend(["", "Limitations:"])
        lines.extend(f"- {item}" for item in limitations[:5])

    return "\n".join(lines)


def _normalize_export_section_text(section_id, content, meta=None):
    sid = _as_text(section_id).strip().lower()
    text = _as_text(content)
    if sid == "action_priorities":
        return _normalize_action_priorities_text(text)
    if sid in ("confidence", "confidence_assessment"):
        if (meta or {}).get("_export_packet_primary"):
            return text
        return _normalize_confidence_text(text, meta)
    return text

# --- BRANDING CONFIG ---

IRON_DISPATCH = {
    "accent": "#b44c3d",
    "accent_dark": "#45382d",
    "text": "#332d28",
    "label": "#7e746a",
    "bg_shade": "#efe6dc",
    "shade_alt": "#e0d2c5",
    "paper": "#f8f2ea",
    "line": "#d2c0b2",
    "line_strong": "#b44c3d",
}

DEEP_WATER = {
    "accent": "#16b7c8",
    "accent_dark": "#263d73",
    "text": "#1f2f44",
    "label": "#708399",
    "bg_shade": "#e7edf3",
    "shade_alt": "#d2dce7",
    "paper": "#f4f7fa",
    "line": "#bfd0de",
    "line_strong": "#16b7c8",
}

BONE_FIELD = {
    "accent": "#395e46",
    "accent_dark": "#234333",
    "text": "#2f2b24",
    "label": "#6f675a",
    "bg_shade": "#ede6d8",
    "shade_alt": "#d8d0bf",
    "paper": "#f6f0e3",
    "line": "#cbbfa9",
    "line_strong": "#4c6e57",
}

THEME_COLORS = {
    "IRON_DISPATCH": IRON_DISPATCH,
    "DEEP_WATER": DEEP_WATER,
    "BONE_FIELD": BONE_FIELD,
    "ARCHITECT": IRON_DISPATCH,
    "NEON_DESERT": BONE_FIELD,
    "CARBON_STEEL": DEEP_WATER,
    "STEEL_RUBY": IRON_DISPATCH,
}

# Fixed semantic colors — DO NOT vary by theme
SEM_BLUE = "#273C75"      # Structural — headers, borders, consensus bar
SEM_GREEN = "#1B4332"     # Verified findings
SEM_RED = "#8B1A1A"       # Flagged/audit findings
SEM_AMBER = "#C8922A"     # Conditional
SEM_MUTED = "#888888"     # Attribution lines
SEM_RULE = "#D8D4CC"      # Thin divider rules

# Agent accent colors — left border per contributor
AGENT_COLORS = {
    "OPENAI": "#273C75",
    "ANTHROPIC": "#8B3A00",
    "GOOGLE": "#2A5C1A",
    "PERPLEXITY": "#4A2A8A",
    "MISTRAL": "#1A4A5C",
}

# --- PDF BRANDING CANVAS ---

def _dark_page_bg(canvas, doc):
    """Optional page painting logic (currently simple white for clarity)."""
    canvas.saveState()
    # If we wanted deep dark bg, we'd do it here. 
    # For 'Architect' theme (the user's latest image), we use white bg.
    canvas.restoreState()

# --- EXPORTERS ---

class ExecutiveMemoExporter:
    """Executive Dossier — Dashboard DNA Standard."""
    @staticmethod
    def generate(intelligence_object, output_dir=None):
        meta, sections, structured, _, _ = _extract_parts(intelligence_object)
        
        filename = f"KORUM-OS_DOSSIER_{_safe_filename_part(meta.get('title'))}_{_timestamp()}.pdf"
        filepath = _output_path(filename, output_dir)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter, 
                                leftMargin=35, rightMargin=35, topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        
        # DNA: Color Palette - locked to Bone & Field across legacy theme ids
        theme_id = _normalize_theme(meta.get('theme', 'BONE_FIELD'))
        tc = THEME_COLORS[theme_id]
        
        ACCENT_HEX = tc["accent"]
        TEXT_HEX = tc["text"]
        LABEL_HEX = tc["label"]
        
        # DNA: Styles - Curated for Executive Scan Patterns
        styles.add(ParagraphStyle('ExecLabel', fontSize=6.5, leading=8, textColor=colors.HexColor(LABEL_HEX), fontName='Helvetica-Bold', letterSpacing=0.5))
        styles.add(ParagraphStyle('ExecSig', fontSize=7, leading=10, textColor=colors.HexColor(LABEL_HEX), alignment=TA_RIGHT, fontName='Helvetica'))
        styles.add(ParagraphStyle('ExecTitle', fontSize=24, leading=28, textColor=colors.HexColor(tc["accent_dark"]), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle('ExecImpact', fontSize=11, leading=16, textColor=colors.HexColor(tc["accent_dark"]), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle('ExecBody', fontSize=8.5, leading=12.5, textColor=colors.HexColor(TEXT_HEX), fontName='Helvetica'))
        styles.add(ParagraphStyle('ExecAudit', fontSize=7, leading=9.5, textColor=colors.HexColor(LABEL_HEX), fontName='Helvetica'))
        styles.add(ParagraphStyle('ExecSectionSubhead', fontSize=8, leading=10, textColor=colors.HexColor(tc["accent_dark"]), fontName='Helvetica-BoldOblique'))
        styles.add(ParagraphStyle('ExecTableHead', fontSize=7, leading=8, textColor=colors.HexColor(tc["accent_dark"]), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle('ExecTableCell', fontSize=7.5, leading=9.5, textColor=colors.HexColor(TEXT_HEX), fontName='Helvetica'))
        styles.add(ParagraphStyle('StatBig', fontSize=20, leading=24, textColor=colors.HexColor(tc["accent_dark"]), fontName='Helvetica-Bold', alignment=TA_CENTER))
        styles.add(ParagraphStyle('StatCaption', fontSize=6.5, leading=8, textColor=colors.HexColor(LABEL_HEX), fontName='Helvetica-Bold', alignment=TA_CENTER, letterSpacing=0.5))
        styles.add(ParagraphStyle('PullQuoteInline', fontSize=8, leading=11, textColor=colors.HexColor(TEXT_HEX), fontName='Helvetica-Oblique'))
        styles.add(ParagraphStyle('Cons', fontSize=9, leading=13, textColor=colors.white, fontName='Helvetica-Oblique'))
        styles.add(ParagraphStyle('ConsScore', fontSize=22, leading=24, textColor=colors.white, fontName='Helvetica-Bold', alignment=TA_RIGHT))

        doc._session_id = _as_text(meta.get('session_id') or meta.get('id') or 'KO-INT-9999').upper()
        contributors = intelligence_object.get("council_contributors") or []
        # divergence_analysis no longer rendered in clean report mode
        mission_ctx = intelligence_object.get("_mission_context") or {}
        client_name = _as_text(mission_ctx.get("client", "")).strip() or "DECISION COMMANDER ALPHA"
        artifacts = _report_artifacts(intelligence_object)
        card_results = intelligence_object.get("_card_results") or {}

        story = []

        # 1. --- STRATEGIC HEADER ---
        workflow_label = _as_text(meta.get('workflow')).upper() or 'STRATEGIC_INTEL'
        doc_title = _as_text(meta.get('title')) or 'Command Node'
        
        # Header Metadata (Top Line)
        top_meta = [
            [[Paragraph(f"INTELLIGENCE DOSSIER | CONFIDENTIAL | {workflow_label}", styles['ExecLabel'])],
             [Paragraph(f"<b>KORUM-OS</b><br/>{client_name}", styles['ExecSig'])]]
        ]
        top_tab = Table(top_meta, colWidths=[360, 180])
        top_tab.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 12)]))
        story.append(top_tab)
        
        # Title
        story.append(Paragraph(doc_title, styles['ExecTitle']))
        story.append(Spacer(1, 15))
        
        # The Blue Underline (Thick Command Line)
        story.append(Table([[[Paragraph("", styles['ExecBody'])]]], colWidths=[540], rowHeights=[2.5], style=[('BACKGROUND', (0,0), (-1,-1), colors.HexColor(ACCENT_HEX))]))
        story.append(Spacer(1, 30))

        # 2. --- EXECUTIVE SUMMARY ---
        summary_text = meta.get("summary") or "Intel synthesis required."
        summary_p = Paragraph(f"<b>{escape(summary_text)}</b>", styles['ExecImpact'])

        top_visual = None
        body_artifacts = list(artifacts)
        if body_artifacts:
            top_visual = body_artifacts.pop(0)

        if top_visual:
            chart_img = _decode_image_data(top_visual.get('imageData'), max_width=200, max_height=180)
            if chart_img:
                visual_col = [
                    Paragraph(_artifact_label(top_visual), styles['StatCaption']),
                    Spacer(1, 5),
                    chart_img,
                ]
            else:
                fallback_note = _artifact_unavailable_note(styles) if _is_raw_svg_artifact(top_visual) else Paragraph(escape(_as_text(top_visual.get('content',''))[:150]).replace("\n","<br/>"), styles['ExecAudit'])
                visual_col = [
                    Paragraph(_artifact_label(top_visual), styles['StatCaption']),
                    Spacer(1, 5),
                    fallback_note
                ]
            summary_tab = Table([[[summary_p], visual_col]], colWidths=[324, 216])
            summary_tab.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (1,0), (1,0), 15),
            ]))
            story.append(summary_tab)
        else:
            story.append(summary_p)
        story.append(Spacer(1, 30))

        # Reorder sections to match canonical output_structure
        section_items_pre = _filter_sections(_reorder_sections(sections or {}), intelligence_object, sections)
        # Chart Standard: assign remaining artifacts to their source nodes for inline placement
        node_artifacts = _assign_artifacts_to_nodes(body_artifacts, section_items_pre, contributors)

        # 3. --- KPI STRIP ---
        key_metrics = structured.get("key_metrics") or []
        if key_metrics:
            value_row = []
            label_row = []
            for km in key_metrics[:3]:
                val = _as_text(km.get('value', '—'))
                lab = _as_text(km.get('metric', '')).upper()
                value_row.append(Paragraph(val, styles['StatBig']))
                label_row.append(Paragraph(lab, styles['StatCaption']))
            
            n_cols = len(value_row) or 1
            col_w = 540 / n_cols
            s_tab = Table([value_row, label_row], colWidths=[col_w] * n_cols, style=[
                ('LINEABOVE', (0,0), (-1,0), 0.5, colors.HexColor(SEM_RULE)),
                ('LINEBELOW', (0,-1), (-1,-1), 0.5, colors.HexColor(SEM_RULE)),
                ('TOPPADDING', (0,0), (-1,-1), 12),
                ('BOTTOMPADDING', (0,0), (-1,-1), 12),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ])
            story.append(s_tab)
            story.append(Spacer(1, 30))

        # 4. --- INTELLIGENCE NODES (FULL-WIDTH PROSE) ---
        section_items = _reorder_sections(sections or {})
        for idx, (sid, content) in enumerate(section_items):
            sec_title = sid.replace("_", " ").upper()
            # Clean report mode: section title only — no node numbers, no provider names
            node_label = sec_title

            # Still resolve provider internally for claim lookups (not displayed)
            prov_name = "KORUM"
            if idx < len(contributors):
                prov_name = _as_text(contributors[idx].get('provider', '')).upper()

            # Flatten structured data (dicts/lists) into readable prose text
            content = _normalize_export_section_text(sid, _flatten_structured_value(content), meta)
            content_blocks = _extract_content_blocks(content)

            prov_key = prov_name.lower()
            prov_data = card_results.get(prov_key) or {}

            if not any(block.get("type") == "table" for block in content_blocks):
                provider_table = _first_table_block(prov_data.get("response", ""))
                if provider_table:
                    content_blocks.append(provider_table)

            # Collect chart images for this node (only those with imageData)
            node_arts = node_artifacts.get(idx, [])
            float_charts = []
            remaining_arts = []
            for art in node_arts:
                chart_img = _decode_image_data(art.get('imageData'), max_width=200, max_height=220)
                if chart_img:
                    float_charts.append((art, chart_img))
                else:
                    remaining_arts.append(art)

            if not content_blocks and not float_charts and not remaining_arts:
                continue

            # Node label + thin rule (always full-width)
            story.append(Paragraph(node_label, styles['ExecLabel']))
            story.append(Spacer(1, 6))
            story.append(Table([[""]],  colWidths=[540], rowHeights=[1],
                style=[('BACKGROUND', (0,0), (-1,-1), colors.HexColor(SEM_RULE))]))
            story.append(Spacer(1, 10))

            # PDF reliability mode: render node prose full-width, then charts immediately after.
            # ReportLab cannot split tall table cells across pages; side-by-side float tables can
            # hard-fail export on long nodes. Stacking preserves proximity without breaking export.
            story.extend(_render_pdf_blocks(content_blocks, styles, 540, tc))
            if float_charts:
                for art, chart_img in float_charts:
                    art_label = _artifact_label(art)
                    if art_label:
                        story.append(Paragraph(art_label, styles['StatCaption']))
                        story.append(Spacer(1, 4))
                    story.append(chart_img)
                    story.append(Spacer(1, 8))

            # Pull quotes removed — clean report mode, no system leakage

            # Remaining artifacts without imageData — render below as text (skip empty)
            for art in remaining_arts:
                if _is_raw_svg_artifact(art):
                    art_label = _artifact_label(art)
                    if art_label:
                        story.append(Paragraph(art_label, styles['ExecSectionSubhead']))
                        story.append(Spacer(1, 4))
                    story.append(_artifact_unavailable_note(styles))
                    story.append(Spacer(1, 10))
                    continue
                art_content = _as_text(art.get('content', '')).strip()
                if not art_content:
                    continue  # skip artifacts with no content and no image — bare labels
                art_label = _artifact_label(art)
                if art_label:
                    story.append(Paragraph(art_label, styles['ExecSectionSubhead']))
                    story.append(Spacer(1, 4))
                art_blocks = _extract_content_blocks(art_content)
                story.extend(_render_pdf_blocks(art_blocks, styles, 540, tc))
                story.append(Spacer(1, 10))

            story.append(Spacer(1, 16))

        # Clean report mode: no ANALYSIS PHASES, no CONFIDENCE SCORE, no closing stamp
        # These are internal system artifacts — not part of the decision document
        story.append(Spacer(1, 20))

        try:
            doc.build(story, onFirstPage=_dark_page_bg, onLaterPages=_dark_page_bg)
        except Exception as e:
            raise RuntimeError(f"ReportLab Build Error: {str(e)}") from e
            
        return filepath

class WordExporter:
    """Executive-Grade Word Report — Premium Dashboard Standard."""
    @staticmethod
    def _hex_to_rgb(hex_str):
        if hex_str.startswith('#'): hex_str = hex_str[1:]
        return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))

    @staticmethod
    def _set_cell_background(cell, fill):
        """Helper to set cell shading in python-docx."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _clear_table_borders(table):
        """Remove all default borders from a python-docx table."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else table._tbl.get_or_add_tblPr()
        tbl_pr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0"/>'
            f'<w:insideH w:val="none" w:sz="0" w:space="0"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0"/>'
            f'</w:tblBorders>'
        ))

    @staticmethod
    def _set_run_style(run, *, size=None, bold=False, italic=False, color=None):
        if size is not None:
            run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))

    @staticmethod
    def _clear_cell(cell):
        cell.text = ""

    @staticmethod
    def _add_paragraph(container, text="", *, size=9, bold=False, italic=False, color=None, align=None):
        p = container.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(_as_text(text))
        WordExporter._set_run_style(run, size=size, bold=bold, italic=italic, color=color)
        return p

    @staticmethod
    def _render_word_table(container, block, theme):
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        if not headers or not rows:
            return None

        table = container.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            WordExporter._set_cell_background(cell, theme["shade_alt"].lstrip('#'))
            p = cell.paragraphs[0]
            run = p.add_run(_as_text(header))
            WordExporter._set_run_style(run, size=8, bold=True, color=theme["accent_dark"])

        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                p = cells[idx].paragraphs[0]
                run = p.add_run(_as_text(value))
                WordExporter._set_run_style(run, size=8.5, color=theme["text"])

        return table

    @staticmethod
    def _render_word_blocks(container, blocks, theme):
        for block in blocks:
            if block["type"] == "paragraph":
                WordExporter._add_paragraph(container, block["text"], size=9, color=theme["text"])
            elif block["type"] == "heading":
                WordExporter._add_paragraph(container, block["text"], size=8, bold=True, italic=True, color=theme["accent_dark"])
            elif block["type"] == "table":
                WordExporter._render_word_table(container, block, theme)
            elif block["type"] == "evidence":
                WordExporter._render_evidence_block(container, block, theme)

    @staticmethod
    def _render_evidence_block(container, block, theme):
        """KORUM-style evidence block: claim strip + status + evidence + provenance."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        block_label = _as_text(block.get("label") or "SOURCE VERIFICATION").upper()
        status = block.get("status", "").upper()
        claim = block.get("claim", "")
        evidence_points = block.get("evidence_points", [])
        sources = block.get("sources", [])
        challenges = block.get("challenges", [])

        # Status-based colors
        if status in ("INACCURATE", "FLAGGED", "CHALLENGED"):
            border_color = "A45A52"
            status_symbol = "\u2717"  # ✗
            status_label = "FLAGGED"
            bg_hex = "FDF6F5"
        elif status in ("ACCURATE", "VERIFIED"):
            border_color = "5B7F5E"
            status_symbol = "\u2713"  # ✓
            status_label = "VERIFIED"
            bg_hex = "F4F8F5"
        else:
            border_color = "4A6A7A"
            status_symbol = "\u25ce"  # ◎
            status_label = "CONDITIONAL"
            bg_hex = "F0F3FA"

        # Fallback for table cells — simplified rendering
        if not hasattr(container, 'add_table'):
            if claim:
                p = WordExporter._add_paragraph(container, claim, size=8, italic=True, color=theme["text"])
                if p:
                    p.paragraph_format.left_indent = Inches(0.3)
            return

        # ── 1. CLAIM STRIP (left-bordered, prominent) ────────────────────
        if claim:
            claim_tab = container.add_table(rows=1, cols=1)
            WordExporter._clear_table_borders(claim_tab)
            claim_cell = claim_tab.rows[0].cells[0]
            tc_pr = claim_cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="18" w:space="0" w:color="{border_color}"/>'
                f'<w:top w:val="none" w:sz="0" w:space="0"/>'
                f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(borders)
            WordExporter._add_paragraph(claim_cell, claim, size=9, bold=True, italic=True, color=theme["text"])
            WordExporter._add_paragraph(claim_cell, f"\u2014 {status_label} CLAIM", size=7, color=border_color)

        # ── 2. STATUS STRIP ──────────────────────────────────────────────
        status_tab = container.add_table(rows=1, cols=1)
        WordExporter._clear_table_borders(status_tab)
        status_cell = status_tab.rows[0].cells[0]
        WordExporter._set_cell_background(status_cell, bg_hex)
        status_text = f"{block_label}   {status_symbol} {status_label}"
        WordExporter._add_paragraph(status_cell, status_text, size=7, bold=True, color=border_color)

        # ── 3. CHALLENGES (if any) ───────────────────────────────────────
        for challenge in challenges[:2]:
            ch_tab = container.add_table(rows=1, cols=1)
            WordExporter._clear_table_borders(ch_tab)
            ch_cell = ch_tab.rows[0].cells[0]
            tc_pr = ch_cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="8" w:space="0" w:color="A45A52"/>'
                f'<w:top w:val="none" w:sz="0" w:space="0"/>'
                f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(borders)
            WordExporter._add_paragraph(ch_cell, challenge, size=8, italic=True, color="A45A52")

        # ── 4. WHY THIS HOLDS (evidence points) ─────────────────────────
        if evidence_points:
            WordExporter._add_paragraph(container, "WHY THIS HOLDS:", size=7, bold=True, color=border_color)
            for point in evidence_points[:5]:
                WordExporter._add_paragraph(container, f"\u2022 {point}", size=8, color=theme["text"])

        # ── 5. PROVENANCE (compact sources) ──────────────────────────────
        if sources:
            source_summary = " \u00b7 ".join(s[:60] for s in sources[:3])
            prov_tab = container.add_table(rows=1, cols=1)
            WordExporter._clear_table_borders(prov_tab)
            prov_cell = prov_tab.rows[0].cells[0]
            WordExporter._set_cell_background(prov_cell, bg_hex)
            WordExporter._add_paragraph(prov_cell, f"SOURCE: {source_summary}", size=6.5, color=SEM_MUTED)
            challenge_status = "PASSED" if status in ("ACCURATE", "VERIFIED") else "FAILED" if status in ("INACCURATE", "FLAGGED") else "PENDING"
            WordExporter._add_paragraph(prov_cell, f"STATUS: {status_symbol} {status_label}   CHALLENGE: {challenge_status}", size=6.5, bold=True, color=border_color)

    @staticmethod
    def _add_spacing(doc, count=1):
        for _ in range(count):
            doc.add_paragraph()

    @staticmethod
    def _ensure_sentence(text):
        clean = _as_text(text).strip()
        if not clean:
            return ""
        if clean[-1] in ".!?":
            return clean
        return clean + "."

    @staticmethod
    def _resolve_final_assessment(meta, sections):
        ordered_sections = dict(_reorder_sections(sections or {}))
        candidates = [
            ordered_sections.get("final_assessment"),
            ordered_sections.get("decision"),
            ordered_sections.get("confidence_assessment"),
            ordered_sections.get("execution_considerations"),
            meta.get("summary"),
        ]
        for candidate in candidates:
            text = _flatten_structured_value(candidate).strip()
            if text:
                return WordExporter._ensure_sentence(text)
        return "KorumOS recommends proceeding on the strongest available decision path while maintaining visible accountability for execution and follow-up review."

    @staticmethod
    def _append_final_assessment_and_footer(doc, meta, sections, theme, intelligence_object):
        session_id = _resolve_session_id(meta, intelligence_object)
        final_assessment = WordExporter._resolve_final_assessment(meta, sections)
        closing_line_1, closing_line_2, closing_line_3 = _closing_stamp_parts(session_id)

        WordExporter._add_spacing(doc)

        fa_label = doc.add_paragraph()
        fa_label_run = fa_label.add_run("FINAL ASSESSMENT")
        WordExporter._set_run_style(
            fa_label_run,
            size=8,
            bold=True,
            color=theme["label"],
        )

        fa_body = doc.add_paragraph()
        fa_body_run = fa_body.add_run(final_assessment)
        WordExporter._set_run_style(
            fa_body_run,
            size=9.5,
            color=theme["text"],
        )

        WordExporter._add_spacing(doc)

        rule_tab = doc.add_table(rows=1, cols=1)
        WordExporter._set_cell_background(rule_tab.rows[0].cells[0], theme["line"].lstrip('#'))
        rule_p = rule_tab.rows[0].cells[0].paragraphs[0]
        rule_p.paragraph_format.space_before = Pt(0)
        rule_p.paragraph_format.space_after = Pt(0)
        rule_p.paragraph_format.line_spacing = Pt(1)
        rule_run = rule_p.add_run("")
        rule_run.font.size = Pt(1)
        from docx.enum.table import WD_ROW_HEIGHT_RULE as _FooterRH
        rule_tab.rows[0].height = Inches(0.03)
        rule_tab.rows[0].height_rule = _FooterRH.EXACTLY

        footer_tab = doc.add_table(rows=2, cols=2)
        footer_tab.columns[0].width = Inches(4.5)
        footer_tab.columns[1].width = Inches(1.8)
        WordExporter._clear_table_borders(footer_tab)

        left_top = footer_tab.rows[0].cells[0]
        left_bottom = footer_tab.rows[1].cells[0]
        right_top = footer_tab.rows[0].cells[1]
        right_bottom = footer_tab.rows[1].cells[1]

        WordExporter._add_paragraph(left_top, closing_line_1, size=8, bold=True, color=theme["accent_dark"])
        WordExporter._add_paragraph(left_bottom, closing_line_2, size=7, color=theme["label"])
        WordExporter._add_paragraph(right_top, closing_line_3, size=7, bold=True, color=theme["accent_dark"], align=WD_ALIGN_PARAGRAPH.RIGHT)
        WordExporter._add_paragraph(right_bottom, "EXPORT STATUS | DECISION ARTIFACT", size=7, bold=True, color=theme["label"], align=WD_ALIGN_PARAGRAPH.RIGHT)

    @staticmethod
    def generate(intelligence_object, output_dir=None):
        meta, sections, structured, _, _ = _extract_parts(intelligence_object)
        doc = Document()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # DNA: SHARED THEME - specialist dynamic resolution
        theme_id = _normalize_theme(meta.get('theme', 'BONE_FIELD'))
        tc = THEME_COLORS[theme_id]
        
        ACCENT_HEX = tc["accent"].lstrip('#')
        DIM_GRAY = tc["label"].lstrip('#')
        s_rgb = RGBColor(*WordExporter._hex_to_rgb(ACCENT_HEX))

        mission_ctx = intelligence_object.get("_mission_context") or {}
        client_name = _as_text(mission_ctx.get("client", "")).strip() or "DECISION COMMANDER ALPHA"
        card_results = intelligence_object.get("_card_results") or {}
        contributors = intelligence_object.get("council_contributors") or []
        # divergence_analysis no longer rendered in clean report mode
        artifacts = _report_artifacts(intelligence_object)

        # Reorder sections to match canonical output_structure
        section_items_pre = [
            (sid, content)
            for sid, content in _filter_sections(_reorder_sections(sections or {}), intelligence_object, sections)
            if _as_text(sid).strip().lower() != "final_assessment"
        ]
        # Chart Standard: assign artifacts to their source nodes for inline placement
        node_artifacts = _assign_artifacts_to_nodes(artifacts, section_items_pre, contributors)

        # 1. --- PREMIUM HEADER ---
        h_tab = doc.add_table(rows=1, cols=2)
        h_tab.columns[0].width = Inches(3.8)
        h_tab.columns[1].width = Inches(2.2)
        
        l_p = h_tab.rows[0].cells[0].paragraphs[0]
        l_run = l_p.add_run(f"INTELLIGENCE DOSSIER  \u00b7  CONFIDENTIAL  \u00b7  {_as_text(meta.get('workflow')).upper() or 'STRATEGIC_INTEL'}")
        l_run.font.size = Pt(6.5)
        l_run.font.bold = True
        l_run.font.color.rgb = RGBColor.from_string(DIM_GRAY)
        
        r_p = h_tab.rows[0].cells[1].paragraphs[0]
        r_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_run = r_p.add_run(f"KORUM-OS\n{client_name}")
        r_run.font.size = Pt(7)
        r_run.font.color.rgb = RGBColor.from_string(DIM_GRAY)
        
        # Title
        t_p = doc.add_paragraph()
        t_run = t_p.add_run(_as_text(meta.get('title') or 'Command Node'))
        t_run.bold = True
        t_run.font.size = Pt(22)
        t_run.font.color.rgb = s_rgb
        
        # Thick Rule (Executive Baseline) — exact height, no cell expansion
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        u_tab = doc.add_table(rows=1, cols=1)
        WordExporter._set_cell_background(u_tab.rows[0].cells[0], ACCENT_HEX)
        # Shrink default paragraph to prevent cell height expansion
        u_p = u_tab.rows[0].cells[0].paragraphs[0]
        u_p.paragraph_format.space_before = Pt(0)
        u_p.paragraph_format.space_after = Pt(0)
        u_p.paragraph_format.line_spacing = Pt(1)
        u_run = u_p.add_run("")
        u_run.font.size = Pt(1)
        u_tab.rows[0].height = Inches(0.04)
        u_tab.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        doc.add_paragraph() # Spacer

        # 2. --- EXECUTIVE SUMMARY ---
        summary_text = meta.get("summary") or "Intel synthesis required."
        s_p = doc.add_paragraph()
        s_run = s_p.add_run(summary_text)
        s_run.bold = True
        s_run.font.size = Pt(11)
        s_run.font.color.rgb = s_rgb

        WordExporter._add_spacing(doc)

        # 3. --- KPI STRIP (thin rules, no box) ---
        key_metrics = structured.get("key_metrics") or []
        if key_metrics:
            kpi_tab = doc.add_table(rows=2, cols=min(3, len(key_metrics)))
            # Thin top/bottom rules only — no full grid box
            from docx.oxml import parse_xml as _px_kpi
            from docx.oxml.ns import nsdecls as _ns_kpi
            from docx.enum.table import WD_ROW_HEIGHT_RULE as _RH
            tbl_pr = kpi_tab._tbl.tblPr if kpi_tab._tbl.tblPr is not None else kpi_tab._tbl.get_or_add_tblPr()
            tbl_borders = _px_kpi(
                f'<w:tblBorders {_ns_kpi("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{tc["line"].lstrip("#")}"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{tc["line"].lstrip("#")}"/>'
                f'<w:left w:val="none" w:sz="0" w:space="0"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0"/>'
                f'<w:insideH w:val="none" w:sz="0" w:space="0"/>'
                f'<w:insideV w:val="none" w:sz="0" w:space="0"/>'
                f'</w:tblBorders>'
            )
            tbl_pr.append(tbl_borders)
            for i, km in enumerate(key_metrics[:3]):
                v_p = kpi_tab.rows[0].cells[i].paragraphs[0]
                v_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                v_run = v_p.add_run(_as_text(km.get('value', '')))
                v_run.bold = True
                v_run.font.size = Pt(18)
                v_run.font.color.rgb = s_rgb
                l_p = kpi_tab.rows[1].cells[i].paragraphs[0]
                l_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                l_run = l_p.add_run(_as_text(km.get('metric', '')).upper())
                l_run.font.size = Pt(7)
                l_run.font.color.rgb = RGBColor.from_string(DIM_GRAY)
            WordExporter._add_spacing(doc)

        # 4. --- INTELLIGENCE NODES (full-width prose) ---
        section_items = section_items_pre
        for idx, (sid, content) in enumerate(section_items):
            sec_title = sid.replace("_", " ").upper()
            # Clean report mode: section title only — no node numbers, no provider names
            prov_name = "KORUM"
            if idx < len(contributors):
                prov_name = _as_text(contributors[idx].get('provider', '')).upper()

            # Flatten structured data (dicts/lists) into readable prose text
            content = _normalize_export_section_text(sid, _flatten_structured_value(content), meta)

            content_blocks = _extract_content_blocks(content)
            prov_key = prov_name.lower()
            prov_data = card_results.get(prov_key) or {}

            if not any(block.get("type") == "table" for block in content_blocks):
                provider_table = _first_table_block(prov_data.get("response", ""))
                if provider_table:
                    content_blocks.append(provider_table)

            # Collect chart images for this node
            node_arts = node_artifacts.get(idx, [])
            float_charts = []
            remaining_arts = []
            for art in node_arts:
                img_stream, img_w, img_h = _decode_image_bytes(art.get('imageData'), max_width_inches=2.4, max_height_inches=2.5)
                if img_stream:
                    float_charts.append((art, img_stream, img_w, img_h))
                else:
                    remaining_arts.append(art)

            if not content_blocks and not float_charts and not remaining_arts:
                continue

            n_p = doc.add_paragraph()
            n_run = n_p.add_run(sec_title)
            n_run.font.size = Pt(7)
            n_run.font.color.rgb = RGBColor.from_string(DIM_GRAY)

            if float_charts:
                # FLOAT MODE: 60/40 side-by-side — prose left, chart right
                float_tab = doc.add_table(rows=1, cols=2)
                float_tab.columns[0].width = Inches(3.6)
                float_tab.columns[1].width = Inches(2.4)
                left_cell = float_tab.rows[0].cells[0]
                right_cell = float_tab.rows[0].cells[1]

                # Left: prose content
                WordExporter._render_word_blocks(left_cell, content_blocks, tc)

                # Right: chart image(s) with labels
                for art, img_stream, img_w, img_h in float_charts:
                    art_label = _artifact_label(art)
                    if art_label:
                        WordExporter._add_paragraph(right_cell, art_label, size=7, bold=True, italic=True, color=tc["accent_dark"])
                    pic_kwargs = {}
                    if img_w:
                        pic_kwargs['width'] = Inches(img_w)
                    if img_h:
                        pic_kwargs['height'] = Inches(img_h)
                    if not pic_kwargs:
                        pic_kwargs['width'] = Inches(2.2)
                    p = right_cell.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_stream, **pic_kwargs)
            else:
                # No chart — full-width prose
                WordExporter._render_word_blocks(doc, content_blocks, tc)

            # Pull quotes removed — clean report mode, no system leakage

            # Remaining text-only artifacts — render below full-width (skip empty)
            for art in remaining_arts:
                if _is_raw_svg_artifact(art):
                    art_label = _artifact_label(art)
                    if art_label:
                        WordExporter._add_paragraph(doc, art_label, size=8, bold=True, italic=True, color=tc["accent_dark"])
                    WordExporter._add_paragraph(doc, "Chart artifact unavailable for export capture.", size=8, italic=True, color=DIM_GRAY)
                    continue
                art_content = _as_text(art.get('content', '')).strip()
                if not art_content:
                    continue  # skip artifacts with no content and no image — bare labels
                art_label = _artifact_label(art)
                if art_label:
                    WordExporter._add_paragraph(doc, art_label, size=8, bold=True, italic=True, color=tc["accent_dark"])
                art_blocks = _extract_content_blocks(art_content)
                WordExporter._render_word_blocks(doc, art_blocks, tc)

            WordExporter._add_spacing(doc)

        # Force a deterministic terminal section so DOCX never ends on an arbitrary node.
        WordExporter._append_final_assessment_and_footer(doc, meta, sections, tc, intelligence_object)

        filename = f"KORUM-OS_DOSSIER_{_safe_filename_part(meta.get('title'))}_{_timestamp()}.docx"
        filepath = _output_path(filename, output_dir)
        tmp_fd, tmp_path = tempfile.mkstemp(prefix="korum_docx_", suffix=".docx", dir=os.path.dirname(filepath))
        os.close(tmp_fd)
        try:
            doc.save(tmp_path)
            with open(tmp_path, 'rb') as handle:
                data = handle.read()
                if not data:
                    raise RuntimeError("DOCX export produced an empty file")
            with zipfile.ZipFile(tmp_path, 'r') as _:
                pass
            os.replace(tmp_path, filepath)
            with open(filepath, 'rb+') as final_handle:
                os.fsync(final_handle.fileno())
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        return filepath

class JSONExporter:
    @staticmethod
    def generate(intelligence_object, output_dir=None):
        filename = f"KORUM-OS_DUMP_{_timestamp()}.json"
        filepath = _output_path(filename, output_dir)
        with open(filepath, 'w') as f:
            json.dump(intelligence_object, f, indent=4)
        return filepath

class CSVExporter:
    @staticmethod
    def generate(intelligence_object, output_dir=None):
        filename = f"KORUM-OS_EXPORT_{_timestamp()}.csv"
        filepath = _output_path(filename, output_dir)
        meta, sections, _, _, _ = _extract_parts(intelligence_object)
        with open(filepath, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Field", "Value"])
            writer.writerow(["Title", meta.get('title')])
            writer.writerow(["Summary", meta.get('summary')])
            for k, v in _reorder_sections(sections or {}):
                writer.writerow([k, _flatten_structured_value(v)])
        return filepath

class TextExporter:
    @staticmethod
    def generate(intelligence_object, output_dir=None):
        meta, sections, structured, _, _ = _extract_parts(intelligence_object)
        lines = [f"KORUM-OS REPORT: {meta.get('title','').upper()}", "="*40, ""]
        lines.append(f"SUMMARY: {meta.get('summary','')}\n")
        for k, v in _reorder_sections(sections or {}):
            lines.append(f"SECTION: {k.upper()}\n{_flatten_structured_value(v)}\n")
        filename = f"KORUM-OS_BRIEF_{_timestamp()}.txt"
        filepath = _output_path(filename, output_dir)
        with open(filepath, 'w') as f:
            f.write("\n".join(lines))
        return filepath

class ResearchPaperWordExporter:
    @staticmethod
    def generate(intelligence_object, output_dir=None):
        return WordExporter.generate(intelligence_object, output_dir)

# Aliases for unified routing
PDFExporter = ExecutiveMemoExporter
MarkdownExporter = TextExporter # Fallback
PPTXExporter = TextExporter     # Placeholder
ResearchPaperExporter = ExecutiveMemoExporter
ExcelExporter = CSVExporter      # Fallback
