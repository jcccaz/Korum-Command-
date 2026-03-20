import tempfile
import unittest
from pathlib import Path

from docx import Document

from exporters import WordExporter


class WordExporterTests(unittest.TestCase):
    def test_docx_export_renders_real_tables_and_normalized_truth(self):
        intelligence_object = {
            "meta": {
                "title": "Bozo Inc Financial Viability",
                "summary": "Revenue remains healthy but cash flow controls need work.",
                "workflow": "FINANCE",
                "theme": "BONE_FIELD",
                "composite_truth_score": 0.85,
                "session_id": "KO-INT-0094",
            },
            "sections": {
                "unit_economics_modeling": (
                    "Base Case Analysis\n"
                    "Monthly Recurring Revenue (MRR): $37,500\n"
                    "Monthly COGS: $15,000\n"
                    "Monthly Operating Burn: $10,000\n"
                    "Monthly Net Cash Generation: $12,500"
                ),
                "sensitivity_table": "Scenario review is below.",
            },
            "structured_data": {
                "key_metrics": [
                    {"metric": "Total Revenue", "value": "$20K"},
                    {"metric": "Net Profit", "value": "$8K"},
                    {"metric": "LTV / CAC", "value": "12.35x"},
                ]
            },
            "council_contributors": [
                {"provider": "openai"},
                {"provider": "perplexity"},
            ],
            "_card_results": {
                "openai": {
                    "verified_claims": [{"claim": "Revenue extraction confirmed across all three streams."}],
                    "response": "No table here.",
                },
                "perplexity": {
                    "verified_claims": [{"claim": "Sensitivity ranges confirmed across modeled scenarios."}],
                    "response": (
                        "| Scenario | Revenue | Net Income |\n"
                        "|---|---:|---:|\n"
                        "| Base | $450,000 | $150,000 |\n"
                        "| Down 10% | $405,000 | $105,000 |\n"
                        "| Up 25% | $562,500 | $262,500 |"
                    ),
                },
            },
            "divergence_analysis": {
                "divergence_summary": "Revenue is verified. Cash controls require remediation."
            },
            "_mission_context": {"client": "BOZO Inc"},
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(WordExporter.generate(intelligence_object, output_dir=tmpdir))
            self.assertTrue(docx_path.exists())

            doc = Document(str(docx_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = [
                [[cell.text.strip() for cell in row.cells] for row in table.rows]
                for table in doc.tables
            ]

            self.assertIn("COUNCIL CONSENSUS: 85/100 TRUTH SCORE", paragraphs)
            self.assertTrue(any("SESSION KO-INT-0094" in " ".join(sum(table, [])) for table in tables))
            self.assertTrue(any(
                any("Monthly Recurring Revenue (MRR)" in cell for row in table for cell in row)
                for table in tables
            ))
            self.assertTrue(any(
                any("Down 10%" in cell for row in table for cell in row)
                for table in tables
            ))


if __name__ == "__main__":
    unittest.main()
